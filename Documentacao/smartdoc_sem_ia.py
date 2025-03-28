"""
Automação para Geração de Documentação de Projetos Power BI
-----------------------------------------------------------

Este código automatiza a documentação de relatórios Power BI a partir de arquivos `.pbit` 
convertidos em `.zip`. Ele processa todos os arquivos .pbit em uma pasta, cria uma subpasta 
para cada um e gera a documentação dentro da respectiva pasta.

Dependências:
- python-docx
- pywin32
"""

import json
import os
import zipfile
import shutil
from datetime import datetime
from docx import Document
import config as cfg

def processar_pbits():
    """Processa todos os arquivos .pbit encontrados na pasta configurada"""
    if not hasattr(cfg, 'arquivos_pbit') or not cfg.arquivos_pbit:
        print("Nenhum arquivo .pbit encontrado para processar!")
        return

    for arquivo_pbit in cfg.arquivos_pbit:
        nome_base = arquivo_pbit.stem
        caminho_pasta_pbit = arquivo_pbit.parent
        
        # Cria pasta para o arquivo (substitui se existir)
        pasta_destino = caminho_pasta_pbit / nome_base
        if pasta_destino.exists():
            # Remove a pasta existente e seu conteúdo
            shutil.rmtree(pasta_destino)
        os.makedirs(pasta_destino)
        
        # Configura os caminhos para este arquivo específico
        arquivo_zip = pasta_destino / f"{nome_base}.zip"
        arquivo_documentacao = pasta_destino / f"Documentação_{nome_base}.docx"
        
        try:
            print(f"\nProcessando arquivo: {arquivo_pbit.name}")
            
            # Processa o arquivo individual
            verificar_ou_renomear_arquivo(arquivo_pbit, arquivo_zip)
            extrair_arquivos_zip(arquivo_zip, pasta_destino)

            # Carrega dados
            layout_data = carregar_dados_json(pasta_destino / 'Report/Layout')
            model_data = carregar_dados_json(pasta_destino / 'DataModelSchema')

            # Reverte para .pbit
            os.rename(arquivo_zip, arquivo_pbit)

            # Extrai informações
            extracoes = {
                "Páginas": extrair_paginas(layout_data),
                "Tabelas": extrair_tabelas(model_data),
                "Medidas": extrair_medidas(model_data),
                "Visuais": extrair_visuais(layout_data),
                "Fontes": extrair_fontes(model_data),
                "Relacionamentos": extrair_relacionamentos(model_data)
            }

            # Gera documento
            gerar_documento(extracoes, arquivo_pbit, arquivo_documentacao)
            print(f"Documentação gerada com sucesso para {arquivo_pbit.name}")
            
        except Exception as e:
            print(f"Erro ao processar {arquivo_pbit.name}: {str(e)}")
            # Limpeza em caso de erro
            if arquivo_zip.exists():
                os.remove(arquivo_zip)

def verificar_ou_renomear_arquivo(arquivo_pbit, arquivo_zip):
    """Verifica e renomeia o arquivo .pbit para .zip se necessário"""
    if arquivo_zip.exists():
        print("Arquivo .zip já existe. Pulando para a próxima instrução.")
    elif arquivo_pbit.exists():
        os.rename(arquivo_pbit, arquivo_zip)
    else:
        raise FileNotFoundError(f"Arquivo não encontrado: {arquivo_pbit}")

def extrair_arquivos_zip(arquivo_zip, pasta_destino):
    """Extrai os arquivos necessários do arquivo ZIP"""
    arquivos_para_extrair = ['Report/Layout', 'DataModelSchema']
    with zipfile.ZipFile(arquivo_zip, 'r') as zip_ref:
        for arquivo in arquivos_para_extrair:
            try:
                zip_ref.extract(arquivo, pasta_destino)
            except KeyError:
                print(f"Aviso: Arquivo {arquivo} não encontrado no ZIP")

def carregar_dados_json(arquivo: str, encoding: str = 'utf-16-le') -> dict:
    """Carrega dados de um arquivo JSON."""
    try:
        with open(arquivo, 'r', encoding=encoding) as f:
            return json.load(f)
    except Exception as e:
        print(f"Erro ao carregar JSON: {arquivo} - {e}")
        return {}

def extrair_paginas(layout: dict) -> str:
    """Extrai informações de páginas."""
    output = []
    for section in layout.get('sections', []):
        page_name = section.get('displayName', 'Sem Nome')
        output.append(f"{page_name}\n-----------\n")
    return "\n".join(output)

def extrair_visuais(layout: dict) -> str:
    """Extrai informações de visuais em cada página."""
    output = []
    for section in layout.get('sections', []):
        page_name = section.get('displayName', 'Sem Nome')
        for container in section.get("visualContainers", []):
            config_data = json.loads(container.get("config", "{}"))
            visual_type = config_data.get("singleVisual", {}).get("visualType")
            position = next(iter(config_data.get("layouts", [])), {}).get("position", {})
            query_refs = [item.get("queryRef") for items in config_data.get("singleVisual", {}).get("projections", {}).values()
                          for item in items if item.get("queryRef")]
            
            output.append(
                f"Página: {page_name}\n"
                f"Posição: X={position.get('x', 0)}, Y={position.get('y', 0)}\n"
                f"Dimensões: {position.get('width', 0)}x{position.get('height', 0)}\n"
                f"Tipo: {visual_type}\n"
                f"Medidas: {', '.join(query_refs) if query_refs else 'Nenhuma'}\n"
                "-----------\n"
            )
    return "\n".join(output)

def extrair_tabelas(model_data: dict) -> str:
    """Extrai informações de tabelas e colunas."""
    output = []
    for table in model_data.get('model', {}).get('tables', []):
        table_name = table.get("name", "")
        if table_name.startswith(("DateTableTemplate", "LocalDateTable")):
            continue
        for column in table.get('columns', []):
            output.append(
                f"Tabela: {table_name}\n"
                f"Coluna: {column.get('name', '')}\n"
                f"Tipo: {column.get('dataType', '')}\n"
                f"Calculada: {'Sim' if column.get('type', '') in ['calculatedTableColumn', 'calculated'] else 'Não'}\n"
                "-----------\n"
            )
    return "\n".join(output)

def extrair_medidas(model_data: dict) -> str:
    """Extrai informações de medidas."""
    output = []
    for table in model_data.get('model', {}).get('tables', []):
        table_name = table.get("name", "")
        for measure in table.get('measures', []):
            expr = measure.get('expression', '')
            if isinstance(expr, list):
                expr = ' '.join(filter(None, expr))
            output.append(
                f"Tabela: {table_name}\n"
                f"Medida: {measure.get('name', '')}\n"
                f"Expressão: {expr}\n"
                "-----------\n"
            )
    return "\n".join(output)

def extrair_fontes(model_data: dict) -> str:
    """Extrai informações sobre fontes de dados."""
    output = []
    for table in model_data.get('model', {}).get('tables', []):
        table_name = table.get("name", "")
        if table_name.startswith(("DateTableTemplate", "LocalDateTable")):
            continue
        for partition in table.get('partitions', []):
            source = partition.get('source', {})
            expr = source.get('expression')
            if isinstance(expr, list):
                expr = ' '.join(filter(None, expr))
            output.append(
                f"Tabela: {table_name}\n"
                f"Modo: {partition.get('mode')}\n"
                f"Tipo: {source.get('type')}\n"
                f"Fonte: {expr}\n"
                "-----------\n"
            )
    return "\n".join(output)

def extrair_relacionamentos(model_data: dict) -> str:
    """Extrai informações de relacionamentos."""
    output = []
    for relation in model_data.get('model', {}).get('relationships', []):
        from_table = relation.get('fromTable')
        to_table = relation.get('toTable')
        if from_table.startswith(("DateTableTemplate", "LocalDateTable")) or \
           to_table.startswith(("DateTableTemplate", "LocalDateTable")):
            continue
        output.append(
            f"De: {from_table}.{relation.get('fromColumn', '')}\n"
            f"Para: {to_table}.{relation.get('toColumn', '')}\n"
            "-----------\n"
        )
    return "\n".join(output)

def salvar_versao(salvar_path):
    """Gerencia versões do arquivo de documentação."""
    if not os.path.exists(salvar_path):
        return salvar_path
    
    base, ext = os.path.splitext(salvar_path)
    versao = 2
    while os.path.exists(f"{base}_v{versao:02}{ext}"):
        versao += 1
    return f"{base}_v{versao:02}{ext}"

def gerar_documento(extracoes, arquivo_pbit, arquivo_documentacao):
    """Gera o documento Word final."""
    modelo_path = cfg.caminho_modelo_word / cfg.nome_modelo_word
    document = Document(modelo_path)

    # Preenche informações básicas
    for para in document.paragraphs:
        if "Data da documentação:" in para.text:
            para.add_run(f" {datetime.now().strftime('%d/%m/%Y')}")
        elif "Nome do Relatório:" in para.text:
            nome_relatorio = arquivo_pbit.stem
            para.add_run(f" {nome_relatorio}")

    # Insere conteúdo
    for titulo, conteudo in extracoes.items():
        for para in document.paragraphs:
            if para.text.strip() == titulo.capitalize():
                novo_paragrafo = document.add_paragraph(conteudo)
                para._element.addnext(novo_paragrafo._element)
                break

    # Salva com controle de versão
    caminho_final = salvar_versao(arquivo_documentacao)
    document.save(caminho_final)
    print(f'Documentação gerada: {caminho_final}')

def main():
    try:
        # Verifica se o modelo Word existe
        modelo_path = cfg.caminho_modelo_word / cfg.nome_modelo_word
        print(f"\nVerificando modelo Word:\nCaminho completo: {modelo_path}")
        
        if not modelo_path.exists():
            raise FileNotFoundError(f"Modelo Word não encontrado: {modelo_path}")
        
        print("Modelo Word encontrado com sucesso!")
        
        # Processa todos os arquivos .pbit encontrados
        processar_pbits()
        
    except Exception as e:
        print(f"Erro durante a execução: {str(e)}")

if __name__ == '__main__':
    main()